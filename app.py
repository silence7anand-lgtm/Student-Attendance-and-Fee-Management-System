from flask import Flask, render_template, request, redirect, url_for, session, g, flash, send_file
import sqlite3
import os
import io
from datetime import date
from functools import wraps
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import pandas as pd
from werkzeug.security import generate_password_hash, check_password_hash

app = Flask(__name__)
app.secret_key = 'dev_key_very_secret_change_in_prod'
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATABASE = os.path.join(BASE_DIR, 'database.db')
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'static', 'uploads', 'profile_pics')
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_db():
    db = getattr(g, '_database', None)
    if db is None:
        db = g._database = sqlite3.connect(DATABASE)
        db.row_factory = sqlite3.Row
    return db

@app.teardown_appcontext
def close_connection(exception):
    db = getattr(g, '_database', None)
    if db is not None:
        db.close()

def init_db():
    with app.app_context():
        db = get_db()
        with app.open_resource('schema.sql', mode='r') as f:
            db.cursor().executescript(f.read())
        db.commit()
        print("Database Initialized")

# Authentication Decorator
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

@app.route('/')
@login_required
def index():
    if session.get('role') == 'student':
        # Fetch Student Dashboard Data
        db = get_db()
        student_id = session.get('user_id') # In login we stored student ID as user_id for students
        
        student = db.execute('SELECT * FROM students WHERE id = ?', (student_id,)).fetchone()
        
        if student is None:
            # Student might have been deleted
            session.clear()
            flash('Student record not found. Please login again.')
            return redirect(url_for('login'))
        
        # Attendance Summary
        attendance_stats = db.execute('''
            SELECT 
                (SELECT COUNT(DISTINCT date) FROM attendance) as total_days,
                COUNT(*) as present_days
            FROM attendance 
            WHERE student_id = ? AND status = 'Present'
        ''', (student_id,)).fetchone()
        
        # Calculate percentage
        total_days = attendance_stats['total_days'] or 0
        present_days = attendance_stats['present_days'] or 0
        percentage = (present_days / total_days * 100) if total_days > 0 else 100
        
        # Fee History
        fees = db.execute('SELECT * FROM fees WHERE student_id = ? ORDER BY date DESC', (student_id,)).fetchall()
        
        # Payment Confirmation Check (Hall Ticket Logic)
        has_paid_regular = len(fees) > 0
        has_paid_condonation = any(f['payment_type'] == 'Condonation' for f in fees)
        
        needs_condonation = percentage < 60
        has_paid = False
        
        if needs_condonation:
            has_paid = has_paid_regular and has_paid_condonation
        else:
            has_paid = has_paid_regular
        
        # Subjects Summary
        subjects = db.execute('SELECT * FROM subjects WHERE class_name = ?', (student['class_name'],)).fetchall()
        
        return render_template('student_dashboard.html', student=student, attendance=attendance_stats, fees=fees, subjects=subjects, has_paid=has_paid, needs_condonation=needs_condonation, has_paid_condonation=has_paid_condonation, percentage=percentage)
        
    return render_template('index.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username_or_roll = request.form['username']
        password = request.form['password']
        
        db = get_db()
        error = None
        user = None
        role = None

        # 1. Try to find user in 'users' table (Admins/Management)
        db_user = db.execute(
            'SELECT * FROM users WHERE username = ?', (username_or_roll,)
        ).fetchone()

        if db_user:
            if check_password_hash(db_user['password_hash'], password):
                user = db_user
                role = db_user['role']
            else:
                error = 'Incorrect password.'
        
        # 2. If not Admin, try to find Student
        if user is None and error is None:
            student_user = db.execute(
                'SELECT * FROM students WHERE roll_no = ?', (username_or_roll,)
            ).fetchone()
            
            if student_user:
                # If no password hash, accept roll_no as default password and set it
                if not student_user['password_hash']:
                    if password == str(student_user['roll_no']):
                        # Auto-set the password hash for future logins
                        db.execute(
                            'UPDATE students SET password_hash = ? WHERE id = ?',
                            (generate_password_hash(password), student_user['id'])
                        )
                        db.commit()
                        user = student_user
                        role = 'student'
                    else:
                        error = 'Incorrect password. Use your Roll Number as password.'
                elif check_password_hash(student_user['password_hash'], password):
                    user = student_user
                    role = 'student'
                else:
                    error = 'Incorrect password. Use your Roll Number as password.'
            else:
                error = 'Incorrect Username/Roll Number or Password.'

        if user and error is None:
            session.clear()
            session['user_id'] = user['id']
            session['role'] = role
            session['name'] = user['name'] if role == 'student' else user['username']
            return redirect(url_for('index'))

        flash(error)

    return render_template('login.html')

@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))

# Placeholder routes for now
@app.route('/register', methods=['GET', 'POST'])
@login_required
def register():
    # Only Admin/Management should be able to register students
    if session.get('role') == 'student':
        flash('Access denied. Admin only.')
        return redirect(url_for('index'))

    if request.method == 'POST':
        name = request.form['name']
        roll_no = request.form['roll_no']
        class_name = request.form['class_name']
        section = request.form.get('section', '')
        contact = request.form['contact']
        
        # Default password is the roll_no
        password_hash = generate_password_hash(roll_no)
        
        db = get_db()
        try:
            db.execute(
                'INSERT INTO students (name, roll_no, class_name, section, contact_info, password_hash) VALUES (?, ?, ?, ?, ?, ?)',
                (name, roll_no, class_name, section, contact, password_hash)
            )
            db.commit()
            flash('Student registered successfully!')
            return redirect(url_for('register'))
        except sqlite3.IntegrityError:
            flash(f'Error: Roll Number {roll_no} already exists.')
            
    # Fetch students to show in a list
    db = get_db()
    students = db.execute('SELECT * FROM students ORDER BY id DESC').fetchall()
    return render_template('register.html', students=students)

@app.route('/attendance', methods=['GET', 'POST'])
@login_required
def attendance():
    if session.get('role') == 'student':
        flash('Access denied. Admin only.')
        return redirect(url_for('index'))

    db = get_db()
    
    # Default date to today
    selected_date = request.args.get('date', date.today().isoformat())
    
    if request.method == 'POST':
        selected_date = request.form['date']
        present_student_ids = request.form.getlist('present')
        
        # Simple Logic: Clear existing attendance for this date and re-insert
        db.execute('DELETE FROM attendance WHERE date = ?', (selected_date,))
        
        for student_id in present_student_ids:
            db.execute(
                'INSERT INTO attendance (student_id, date, status) VALUES (?, ?, ?)',
                (student_id, selected_date, 'Present')
            )
        db.commit()
        flash(f'Attendance updated for {selected_date}')
        return redirect(url_for('attendance', date=selected_date))
        
    # Get all students
    students = db.execute('SELECT * FROM students').fetchall()
    
    # Get attendance for selected date
    attendance_records = db.execute(
        'SELECT student_id FROM attendance WHERE date = ? AND status = "Present"', 
        (selected_date,)
    ).fetchall()
    present_ids = [row['student_id'] for row in attendance_records]
    
    return render_template('attendance.html', students=students, selected_date=selected_date, present_ids=present_ids)

@app.route('/fees')
@login_required
def fees():
    if session.get('role') == 'student':
        flash('Access denied. Admin only.')
        return redirect(url_for('index'))
        
    db = get_db()
    
    # Only show history in admin login (Pay List)
    history = db.execute('''
        SELECT f.*, s.name, s.roll_no 
        FROM fees f 
        JOIN students s ON f.student_id = s.id 
        ORDER BY f.date DESC
    ''').fetchall()
    
    return render_template('fees.html', history=history)

@app.route('/delete_fee/<int:id>', methods=['POST'])
@login_required
def delete_fee(id):
    db = get_db()
    db.execute('DELETE FROM fees WHERE id = ?', (id,))
    db.commit()
    flash('Fee record deleted successfully.')
    return redirect(url_for('fees'))

@app.route('/delete_student/<int:id>', methods=['POST'])
@login_required
def delete_student(id):
    db = get_db()
    # Delete related data first
    db.execute('DELETE FROM attendance WHERE student_id = ?', (id,))
    db.execute('DELETE FROM fees WHERE student_id = ?', (id,))
    # Delete student
    db.execute('DELETE FROM students WHERE id = ?', (id,))
    db.commit()
    flash('Student and related records deleted successfully.')
    return redirect(url_for('register'))

@app.route('/delete_attendance_item', methods=['POST'])
@login_required
def delete_attendance_item():
    student_id = request.form['student_id']
    date = request.form['date']
    db = get_db()
    db.execute('DELETE FROM attendance WHERE student_id = ? AND date = ?', (student_id, date))
    db.commit()
    flash('Attendance record removed.')
    return redirect(url_for('attendance', date=date))


@app.route('/import', methods=['POST'])
@login_required
def import_data():
    if session.get('role') == 'student':
        flash('Access denied. Admin only.')
        return redirect(url_for('index'))

    if 'file' not in request.files:
        flash('No file part')
        return redirect(url_for('register'))
        
    file = request.files['file']
    if file.filename == '':
        flash('No selected file')
        return redirect(url_for('register'))
        
    if file and file.filename.endswith('.xlsx'):
        try:
            df = pd.read_excel(file)
            db = get_db()
            
            count_students = 0
            count_fees = 0
            
            for _, row in df.iterrows():
                # 1. Provide default column names if they match loosely
                name = row.get('Name')
                roll_no = row.get('RollNo')
                class_name = row.get('Class')
                section = row.get('Section', '')
                contact = row.get('Contact')
                
                if name and roll_no:
                    # Check if student exists
                    exist = db.execute('SELECT id FROM students WHERE roll_no = ?', (roll_no,)).fetchone()
                    student_id = None
                    
                    if not exist:
                        # Default password is the roll_no
                        password_hash = generate_password_hash(str(roll_no))
                        cursor = db.execute(
                            'INSERT INTO students (name, roll_no, class_name, section, contact_info, password_hash) VALUES (?, ?, ?, ?, ?, ?)',
                            (name, roll_no, class_name, section, contact, password_hash)
                        )
                        student_id = cursor.lastrowid
                        count_students += 1
                    else:
                        student_id = exist['id']
                    
                    # 2. Process Fees if columns exist
                    amount = row.get('FeeAmount', 0)
                    if amount > 0:
                        late_fee = row.get('LateFee', 0)
                        payment_type = row.get('FeeType', 'Other')
                        remarks = row.get('FeeRemarks', 'Imported')
                        
                        db.execute(
                            'INSERT INTO fees (student_id, amount, late_fee, date, payment_type, remarks) VALUES (?, ?, ?, ?, ?, ?)',
                            (student_id, amount, late_fee, date.today().isoformat(), payment_type, remarks)
                        )
                        count_fees += 1

            db.commit()
            flash(f'Imported {count_students} new students and {count_fees} fee records.')
        except Exception as e:
            flash(f'Error processing file: {e}')
            
    return redirect(url_for('register'))

@app.route('/export_students')
@login_required
def export_students():
    if session.get('role') == 'student':
        flash('Access denied. Admin only.')
        return redirect(url_for('index'))

    db = get_db()
    students = db.execute('SELECT * FROM students').fetchall()
    
    # Convert to list of dicts
    data = []
    for s in students:
        data.append({
            'Name': s['name'],
            'RollNo': s['roll_no'],
            'Class': s['class_name'],
            'Section': s['section'],
            'Contact': s['contact_info']
        })
        
    df = pd.DataFrame(data)
    
    # Save to memory buffer
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Students')
        _style_excel_worksheet(writer.sheets['Students'], df)
    output.seek(0)
    
    return send_file(output, download_name='stu.xlsx', as_attachment=True)

@app.route('/export_fees')
@login_required
def export_fees():
    if session.get('role') == 'student':
        flash('Access denied. Admin only.')
        return redirect(url_for('index'))

    db = get_db()
    # Join with students to get names
    fees = db.execute('''
        SELECT f.*, s.name, s.roll_no 
        FROM fees f 
        JOIN students s ON f.student_id = s.id
    ''').fetchall()
    
    data = []
    for f in fees:
        data.append({
            'Date': f['date'],
            'Student Name': f['name'],
            'Roll No': f['roll_no'],
            'Amount': f['amount'],
            'Late Fee': f['late_fee'],
            'Type': f['payment_type'],
            'Remarks': f['remarks']
        })
        
    df = pd.DataFrame(data)
    
    import io
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Fees')
        _style_excel_worksheet(writer.sheets['Fees'], df)
    output.seek(0)
    
    return send_file(output, download_name='fees.xlsx', as_attachment=True)

@app.route('/student_import', methods=['POST'])
@login_required
def student_import():
    if session.get('role') != 'student':
        flash('Access denied. Students only.')
        return redirect(url_for('index'))

    student_id = session.get('user_id')
    
    if 'file' not in request.files:
        flash('No file part')
        return redirect(url_for('index'))
        
    file = request.files['file']
    if file.filename == '':
        flash('No selected file')
        return redirect(url_for('index'))
        
    if file and file.filename.endswith('.xlsx'):
        try:
            df = pd.read_excel(file)
            db = get_db()
            
            count_fees = 0
            
            for _, row in df.iterrows():
                amount = row.get('FeeAmount', 0)
                if amount > 0:
                    late_fee = row.get('LateFee', 0)
                    payment_type = row.get('FeeType', 'Other')
                    remarks = row.get('FeeRemarks', 'Student Import')
                    fee_date = row.get('Date', pd.Timestamp.now().strftime('%Y-%m-%d'))
                    
                    if isinstance(fee_date, pd.Timestamp):
                        fee_date = fee_date.strftime('%Y-%m-%d')
                    
                    db.execute(
                        'INSERT INTO fees (student_id, amount, late_fee, date, payment_type, remarks) VALUES (?, ?, ?, ?, ?, ?)',
                        (student_id, amount, late_fee, fee_date, payment_type, remarks)
                    )
                    count_fees += 1

            db.commit()
            flash(f'Imported {count_fees} fee records to your profile.')
        except Exception as e:
            flash(f'Error processing file: {e}')
            
    return redirect(url_for('student_fees'))

@app.route('/student_fees', methods=['GET', 'POST'])
@login_required
def student_fees():
    if session.get('role') != 'student':
        flash('Access denied. Students only.')
        return redirect(url_for('index'))
        
    db = get_db()
    student_id = session.get('user_id')

    if request.method == 'POST':
        amount = request.form['amount']
        late_fee = request.form.get('late_fee', 0)
        late_fee = float(late_fee) if late_fee else 0
        date = request.form['date']
        payment_type = request.form['payment_type']
        remarks = request.form['remarks']
        
        db.execute(
            'INSERT INTO fees (student_id, amount, late_fee, date, payment_type, remarks) VALUES (?, ?, ?, ?, ?, ?)',
            (student_id, amount, late_fee, date, payment_type, remarks)
        )
        db.commit()
        flash('Payment recorded successfully!')
        return redirect(url_for('student_fees'))
    
    fees = db.execute('SELECT * FROM fees WHERE student_id = ? ORDER BY date DESC', (student_id,)).fetchall()
    
    return render_template('student_fees.html', fees=fees)

@app.route('/student_subjects')
@login_required
def student_subjects():
    if session.get('role') != 'student':
        flash('Access denied. Students only.')
        return redirect(url_for('index'))
        
    db = get_db()
    student_id = session.get('user_id')
    
    # Get student's class
    student = db.execute('SELECT class_name FROM students WHERE id = ?', (student_id,)).fetchone()
    
    if student:
        # Get subjects for that class
        subjects = db.execute('SELECT * FROM subjects WHERE class_name = ?', (student['class_name'],)).fetchall()
    else:
        subjects = []
        
    return render_template('student_subjects.html', subjects=subjects)



# Admin Subject Management
@app.route('/admin/subjects', methods=['GET', 'POST'])
@login_required
def admin_subjects():
    if session.get('role') == 'student':
        flash('Access denied. Admin only.')
        return redirect(url_for('index'))
    
    db = get_db()
    
    if request.method == 'POST':
        name = request.form['name']
        class_name = request.form['class_name']
        teacher = request.form['teacher']
        desc = request.form['description']
        
        db.execute(
            'INSERT INTO subjects (name, class_name, teacher_name, description) VALUES (?, ?, ?, ?)',
            (name, class_name, teacher, desc)
        )
        db.commit()
        flash('Subject added successfully!')
        return redirect(url_for('admin_subjects'))
        
    subjects = db.execute('SELECT * FROM subjects ORDER BY class_name, name').fetchall()
    return render_template('admin_subjects.html', subjects=subjects)

@app.route('/admin/delete_subject/<int:id>', methods=['POST'])
@login_required
def delete_subject(id):
    if session.get('role') == 'student':
        flash('Access denied. Admin only.')
        return redirect(url_for('index'))
        
    db = get_db()
    db.execute('DELETE FROM subjects WHERE id = ?', (id,))
    db.commit()
    flash('Subject deleted successfully.')
    return redirect(url_for('admin_subjects'))

@app.route('/student/<int:id>')
@login_required
def student_detail(id):
    if session.get('role') == 'student':
        flash('Access denied. Admin only.')
        return redirect(url_for('index'))
        
    db = get_db()
    
    # 1. Fetch Student Profile
    student = db.execute('SELECT * FROM students WHERE id = ?', (id,)).fetchone()
    if not student:
        flash('Student record not found.')
        return redirect(url_for('register'))
        
    # 2. Attendance Summary
    attendance_stats = db.execute('''
        SELECT 
            COUNT(*) as total_days,
            SUM(CASE WHEN status = 'Present' THEN 1 ELSE 0 END) as present_days
        FROM attendance 
        WHERE student_id = ?
    ''', (id,)).fetchone()
    
    # 3. Fee History
    fees = db.execute('SELECT * FROM fees WHERE student_id = ? ORDER BY date DESC', (id,)).fetchall()
    
    return render_template('student_detail.html', student=student, attendance=attendance_stats, fees=fees)

@app.route('/upload_profile_pic/<int:student_id>', methods=['POST'])
@login_required
def upload_profile_pic(student_id):
    if session.get('role') == 'student' and session.get('user_id') != student_id:
        flash('Access denied.')
        return redirect(url_for('index'))

    if 'profile_pic' not in request.files:
        flash('No file part')
        return redirect(url_for('student_detail', id=student_id))
    
    file = request.files['profile_pic']
    if file.filename == '':
        flash('No selected file')
        return redirect(url_for('student_detail', id=student_id))
    
    if file and allowed_file(file.filename):
        import uuid
        filename = f"{student_id}_{uuid.uuid4().hex}.{file.filename.rsplit('.', 1)[1].lower()}"
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        db = get_db()
        # Delete old file if exists
        old_pic = db.execute('SELECT profile_pic FROM students WHERE id = ?', (student_id,)).fetchone()
        if old_pic and old_pic['profile_pic']:
            old_path = os.path.join(app.config['UPLOAD_FOLDER'], old_pic['profile_pic'])
            if os.path.exists(old_path):
                try:
                    os.remove(old_path)
                except:
                    pass
        
        db.execute('UPDATE students SET profile_pic = ? WHERE id = ?', (filename, student_id))
        db.commit()
        flash('Profile picture updated successfully!')
    else:
        flash('Invalid file type. Allowed: png, jpg, jpeg, gif')
        
    return redirect(url_for('student_detail', id=student_id))

# ============================================================
# HELPER: Build a professionally styled student report document
# ============================================================
def _style_excel_worksheet(worksheet, df):
    """Apply professional Oxford Blue theme to openpyxl formatting."""
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    
    header_fill = PatternFill(start_color='002147', end_color='002147', fill_type='solid')
    header_font = Font(name='Calibri', color='FFFFFF', bold=True, size=11)
    center_align = Alignment(horizontal='center', vertical='center')
    thin_border = Border(left=Side(style='thin', color='D3D3D3'), 
                         right=Side(style='thin', color='D3D3D3'), 
                         top=Side(style='thin', color='D3D3D3'), 
                         bottom=Side(style='thin', color='D3D3D3'))

    worksheet.row_dimensions[1].height = 20

    col_headers = df.columns if not df.empty else []
    for col_idx, col_name in enumerate(col_headers, 1):
        cell_obj = worksheet.cell(row=1, column=col_idx)
        col_letter = cell_obj.column_letter
        
        try:
            max_len = max(
                len(str(col_name)),
                df[col_name].astype(str).str.len().max() if not df.empty else 0
            )
        except Exception:
            max_len = len(str(col_name))
            
        worksheet.column_dimensions[col_letter].width = min(max_len + 3, 35)
        
        # Style header cell uniquely
        cell_obj.fill = header_fill
        cell_obj.font = header_font
        cell_obj.alignment = center_align
        cell_obj.border = thin_border

def _style_table_header(cell, text):
    """Style a table header cell with Oxford Blue background and white text."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(255, 255, 255)
    # Set cell background color (Oxford Blue)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), '002147')
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)

def _style_table_cell(cell, text, bold=False, color=None):
    """Style a regular table cell."""
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size = Pt(9)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color

def _build_student_report_doc(student, attendance, fees, subjects=None):
    """Build a professionally styled Word document report for a student."""
    document = Document()
    
    # -- Set default font --
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0, 0, 0)
    
    # -- Letterhead --
    logo_path = os.path.join(BASE_DIR, 'static', 'logo.jpg')
    if os.path.exists(logo_path):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(logo_path, width=Inches(0.8))
    
    header_title = document.add_heading('', level=0)
    header_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_title.add_run('BHARATHIDASAN UNIVERSITY')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(26, 54, 93)  # BDU Navy
    run.bold = True
    
    address = document.add_paragraph()
    address.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = address.add_run("Tiruchirappalli, Tamil Nadu, India\nEstablished by the Government of Tamil Nadu")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    # Divider line
    div = document.add_paragraph()
    div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = div.add_run('━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━')
    run.font.color.rgb = RGBColor(200, 152, 44) # Gold
    run.font.size = Pt(9)
    
    # Document Title
    doc_title = document.add_heading('', level=1)
    doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = doc_title.add_run('OFFICIAL STUDENT REPORT')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = True
    
    # -- Student Name --
    name_p = document.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_p.add_run(student['name'])
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(123, 26, 44)
    run.bold = True
    
    # -- Profile Picture (optional) --
    if student['profile_pic']:
        pic_path = os.path.join(UPLOAD_FOLDER, student['profile_pic'])
        if os.path.exists(pic_path):
            try:
                # Add picture centered
                p = document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(pic_path, width=Inches(1.5))
            except Exception as e:
                print(f"Error adding picture to report: {e}")
    
    # -- Profile Table --
    heading = document.add_heading('', level=1)
    run = heading.add_run('Student Profile')
    run.font.color.rgb = RGBColor(123, 26, 44)
    run.font.size = Pt(14)
    
    profile_table = document.add_table(rows=4, cols=2)
    profile_table.style = 'Table Grid'
    profile_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    profile_data = [
        ('Roll Number', student['roll_no']),
        ('Class', student['class_name']),
        ('Contact', student['contact_info'] or 'N/A'),
        ('Report Date', pd.Timestamp.now().strftime('%d %B %Y')),
    ]
    for i, (label, value) in enumerate(profile_data):
        _style_table_header(profile_table.rows[i].cells[0], label)
        _style_table_cell(profile_table.rows[i].cells[1], value, bold=True)
    
    # Set column widths
    for row in profile_table.rows:
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(10)
    
    document.add_paragraph('')  # Spacer
    
    # -- Attendance Summary --
    heading = document.add_heading('', level=1)
    run = heading.add_run('Attendance Summary')
    run.font.color.rgb = RGBColor(123, 26, 44)
    run.font.size = Pt(14)
    
    total_records = len(attendance)
    present_count = sum(1 for r in attendance if r['status'] == 'Present')
    absent_count = total_records - present_count
    percentage = round((present_count / total_records * 100), 1) if total_records > 0 else 0
    
    # Summary stats table
    stats_table = document.add_table(rows=1, cols=4)
    stats_table.style = 'Table Grid'
    stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    stat_headers = ['Total Days', 'Present', 'Absent', 'Percentage']
    stat_values = [str(total_records), str(present_count), str(absent_count), f'{percentage}%']
    stat_colors = [
        RGBColor(45, 27, 27),
        RGBColor(22, 118, 61),
        RGBColor(185, 28, 47),
        RGBColor(22, 118, 61) if percentage >= 75 else RGBColor(185, 28, 47)
    ]
    
    for i, header in enumerate(stat_headers):
        _style_table_header(stats_table.rows[0].cells[i], header)
    
    value_row = stats_table.add_row()
    for i, (val, clr) in enumerate(zip(stat_values, stat_colors)):
        cell = value_row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(val)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = clr
    
    document.add_paragraph('')  # Spacer
    
    # Attendance detail table
    if attendance:
        heading = document.add_heading('', level=2)
        run = heading.add_run('Attendance Records')
        run.font.color.rgb = RGBColor(123, 26, 44)
        run.font.size = Pt(12)
        
        att_table = document.add_table(rows=1, cols=3)
        att_table.style = 'Table Grid'
        att_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for i, h in enumerate(['#', 'Date', 'Status']):
            _style_table_header(att_table.rows[0].cells[i], h)
        
        for idx, row in enumerate(attendance, 1):
            row_cells = att_table.add_row().cells
            _style_table_cell(row_cells[0], str(idx))
            _style_table_cell(row_cells[1], row['date'])
            status_color = RGBColor(22, 118, 61) if row['status'] == 'Present' else RGBColor(185, 28, 47)
            _style_table_cell(row_cells[2], row['status'], bold=True, color=status_color)
    else:
        p = document.add_paragraph()
        run = p.add_run('No attendance records found.')
        run.font.color.rgb = RGBColor(107, 83, 83)
        run.italic = True
    
    document.add_paragraph('')  # Spacer
    
    # -- Fee History --
    heading = document.add_heading('', level=1)
    run = heading.add_run('Fee Payment History')
    run.font.color.rgb = RGBColor(123, 26, 44)
    run.font.size = Pt(14)
    
    if fees:
        # Fee summary
        total_fees = sum(f['amount'] for f in fees)
        total_late = sum(f['late_fee'] for f in fees)
        
        sum_p = document.add_paragraph()
        run = sum_p.add_run(f'Total Paid: Rs. {total_fees:,.2f}')
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(22, 118, 61)
        if total_late > 0:
            run = sum_p.add_run(f'  |  Late Fees: Rs. {total_late:,.2f}')
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(185, 28, 47)
        
        document.add_paragraph('')
        
        fee_table = document.add_table(rows=1, cols=6)
        fee_table.style = 'Table Grid'
        fee_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for i, h in enumerate(['#', 'Date', 'Amount', 'Late Fee', 'Type', 'Remarks']):
            _style_table_header(fee_table.rows[0].cells[i], h)
        
        for idx, row in enumerate(fees, 1):
            row_cells = fee_table.add_row().cells
            _style_table_cell(row_cells[0], str(idx))
            _style_table_cell(row_cells[1], row['date'])
            _style_table_cell(row_cells[2], f"Rs. {row['amount']}", bold=True, color=RGBColor(22, 118, 61))
            late = row['late_fee']
            _style_table_cell(row_cells[3], f"Rs. {late}" if late > 0 else '-', 
                            color=RGBColor(185, 28, 47) if late > 0 else None)
            _style_table_cell(row_cells[4], row['payment_type'] or '-')
            _style_table_cell(row_cells[5], row['remarks'] or '-')
    else:
        p = document.add_paragraph()
        run = p.add_run('No fee records found.')
        run.font.color.rgb = RGBColor(107, 83, 83)
        run.italic = True
    
    # -- Footer --
    document.add_paragraph('')
    footer_div = document.add_paragraph()
    footer_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_div.add_run('━' * 60)
    run.font.color.rgb = RGBColor(200, 152, 44)
    run.font.size = Pt(8)
    
    footer = document.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"Generated on {pd.Timestamp.now().strftime('%d %B %Y at %I:%M %p')} — Campus Portal")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(154, 116, 32)
    run.italic = True
    
    return document


@app.route('/download_report/<int:id>')
@login_required
def download_report(id):
    if session.get('role') == 'student':
        flash('Access denied. Admin only.')
        return redirect(url_for('index'))
        
    student_id = id
    db = get_db()
    
    student = db.execute('SELECT * FROM students WHERE id = ?', (student_id,)).fetchone()
    if not student:
        flash('Student record not found.')
        return redirect(url_for('register'))
        
    attendance = db.execute('SELECT date, status FROM attendance WHERE student_id = ? ORDER BY date DESC', (student_id,)).fetchall()
    fees = db.execute('SELECT date, amount, late_fee, payment_type, remarks FROM fees WHERE student_id = ? ORDER BY date DESC', (student_id,)).fetchall()
    
    document = _build_student_report_doc(student, attendance, fees)
    
    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    
    timestamp = pd.Timestamp.now().strftime('%Y%m%d')
    filename = f"Report_{student['roll_no']}_{timestamp}.docx"
    
    return send_file(output, download_name=filename, as_attachment=True)

@app.route('/download_my_report')
@login_required
def download_my_report():
    if session.get('role') != 'student':
        flash('Access denied. Students only.')
        return redirect(url_for('index'))
        
    student_id = session.get('user_id')
    db = get_db()
    
    student = db.execute('SELECT * FROM students WHERE id = ?', (student_id,)).fetchone()
    if not student:
        flash('Student record not found.')
        return redirect(url_for('index'))
    
    attendance = db.execute('SELECT date, status FROM attendance WHERE student_id = ? ORDER BY date DESC', (student_id,)).fetchall()
    fees = db.execute('SELECT date, amount, late_fee, payment_type, remarks FROM fees WHERE student_id = ? ORDER BY date DESC', (student_id,)).fetchall()
    
    document = _build_student_report_doc(student, attendance, fees)
    
    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    
    timestamp = pd.Timestamp.now().strftime('%Y%m%d')
    filename = f"Report_{student['roll_no']}_{timestamp}.docx"
    
    return send_file(output, download_name=filename, as_attachment=True)


# ============================================================
# ATTENDANCE EXCEL EXPORT
# ============================================================
@app.route('/export_attendance')
@login_required
def export_attendance():
    if session.get('role') == 'student':
        flash('Access denied. Admin only.')
        return redirect(url_for('index'))
    
    db = get_db()
    
    # Get all students
    students = db.execute('SELECT id, name, roll_no, class_name FROM students ORDER BY roll_no').fetchall()
    
    # Get all unique attendance dates
    dates_rows = db.execute('SELECT DISTINCT date FROM attendance ORDER BY date DESC').fetchall()
    all_dates = [r['date'] for r in dates_rows]
    
    # Get all attendance records
    attendance_records = db.execute('''
        SELECT student_id, date, status FROM attendance ORDER BY date
    ''').fetchall()
    
    # Build lookup: (student_id, date) -> status
    att_lookup = {}
    for rec in attendance_records:
        att_lookup[(rec['student_id'], rec['date'])] = rec['status']
    
    # Build data for Excel
    rows = []
    for s in students:
        row = {
            'Roll No': s['roll_no'],
            'Name': s['name'],
            'Class': s['class_name'],
            'Section': s.get('section', ''),
        }
        present_count = 0
        for d in all_dates:
            status = att_lookup.get((s['id'], d), 'Absent')
            row[d] = status
            if status == 'Present':
                present_count += 1
        
        total = len(all_dates) if all_dates else 0
        row['Total Present'] = present_count
        row['Total Absent'] = total - present_count
        row['Percentage'] = f"{round(present_count / total * 100, 1)}%" if total > 0 else "N/A"
        rows.append(row)
    
    df = pd.DataFrame(rows)
    
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Attendance Report')
        _style_excel_worksheet(writer.sheets['Attendance Report'], df)
    
    output.seek(0)
    
    timestamp = pd.Timestamp.now().strftime('%Y%m%d')
    return send_file(output, download_name=f'Attendance_Report_{timestamp}.xlsx', as_attachment=True)

@app.route('/download_hall_ticket')
@login_required
def download_hall_ticket():
    if session.get('role') != 'student':
        flash('Access denied. Students only.')
        return redirect(url_for('index'))
        
    student_id = session.get('user_id')
    db = get_db()
    
    # 1. Fetch Student Profile
    student = db.execute('SELECT * FROM students WHERE id = ?', (student_id,)).fetchone()
    if not student:
        flash('Student record not found.')
        return redirect(url_for('index'))
    
    # 2. Check Payment Status
    fees = db.execute('SELECT * FROM fees WHERE student_id = ?', (student_id,)).fetchall()
    if not fees:
        flash('Hall ticket is available only after fee payment confirmation.')
        return redirect(url_for('index'))
        
    # Check 60% rule
    attendance_stats = db.execute('''
        SELECT 
            (SELECT COUNT(DISTINCT date) FROM attendance) as total_days,
            COUNT(*) as present_days
        FROM attendance 
        WHERE student_id = ? AND status = 'Present'
    ''', (student_id,)).fetchone()
    
    total_days = attendance_stats['total_days'] or 0
    present_days = attendance_stats['present_days'] or 0
    percentage = (present_days / total_days * 100) if total_days > 0 else 100
    
    has_paid_condonation = any(f['payment_type'] == 'Condonation' for f in fees)
    if percentage < 60 and not has_paid_condonation:
        flash('Attendance below 60%. Condonation fee payment is required.')
        return redirect(url_for('index'))
        
    # 3. Fetch Subjects
    subjects = db.execute('SELECT * FROM subjects WHERE class_name = ?', (student['class_name'],)).fetchall()
    
    # Create Word Document
    document = Document()
    
    # -- Set default font --
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0, 0, 0)
    
    # -- Letterhead --
    logo_path = os.path.join(BASE_DIR, 'static', 'logo.jpg')
    if os.path.exists(logo_path):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(logo_path, width=Inches(0.8))

    header_title = document.add_heading('', level=0)
    header_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_title.add_run('BHARATHIDASAN UNIVERSITY')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(26, 54, 93)  # BDU Navy
    run.bold = True
    
    address = document.add_paragraph()
    address.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = address.add_run("Tiruchirappalli, Tamil Nadu, India\nEstablished by the Government of Tamil Nadu")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    # Divider line
    div = document.add_paragraph()
    div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = div.add_run('━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━')
    run.font.color.rgb = RGBColor(200, 152, 44) # Gold
    run.font.size = Pt(9)
    
    # Document Title
    doc_title = document.add_heading('', level=1)
    doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = doc_title.add_run('EXAMINATION HALL TICKET')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = True
    
    # Student Details
    heading = document.add_heading('', level=1)
    run = heading.add_run('Student Details')
    run.font.color.rgb = RGBColor(123, 26, 44)
    run.font.size = Pt(14)
    
    profile_table = document.add_table(rows=3, cols=2)
    profile_table.style = 'Table Grid'
    profile_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    profile_data = [
        ('Name', student['name']),
        ('Roll Number', student['roll_no']),
        ('Class', student['class_name']),
    ]
    for i, (label, value) in enumerate(profile_data):
        _style_table_header(profile_table.rows[i].cells[0], label)
        _style_table_cell(profile_table.rows[i].cells[1], value, bold=True)
    
    for row in profile_table.rows:
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(10)
        
    document.add_paragraph('')  # Spacer
    
    # Subjects List
    heading = document.add_heading('', level=1)
    run = heading.add_run('Examination Subjects')
    run.font.color.rgb = RGBColor(123, 26, 44)
    run.font.size = Pt(14)
    
    if subjects:
        table = document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for i, h in enumerate(['Subject Name', 'Teacher']):
            _style_table_header(table.rows[0].cells[i], h)
            
        for sub in subjects:
            row_cells = table.add_row().cells
            _style_table_cell(row_cells[0], sub['name'], bold=True)
            _style_table_cell(row_cells[1], sub['teacher_name'] or 'N/A')
    else:
        p = document.add_paragraph()
        run = p.add_run('No subjects registered for this class.')
        run.font.color.rgb = RGBColor(107, 83, 83)
        run.italic = True
        
    document.add_paragraph('')
    
    note_p = document.add_paragraph()
    run = note_p.add_run("Note: Please bring this hall ticket to the examination hall along with your ID card.")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(185, 28, 47)
    
    # Save to buffer
    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    
    filename = f"HallTicket_{student['roll_no']}.docx"
    return send_file(output, download_name=filename, as_attachment=True)


@app.route('/download_hall_ticket_pdf')
@login_required
def download_hall_ticket_pdf():
    if session.get('role') != 'student':
        flash('Access denied. Students only.')
        return redirect(url_for('index'))
        
    student_id = session.get('user_id')
    db = get_db()
    
    student = db.execute('SELECT * FROM students WHERE id = ?', (student_id,)).fetchone()
    if not student:
        flash('Student record not found.')
        return redirect(url_for('index'))
    
    # 2. Check Payment Status
    fees = db.execute('SELECT * FROM fees WHERE student_id = ?', (student_id,)).fetchall()
    if not fees:
        flash('Hall ticket is available only after fee payment confirmation.')
        return redirect(url_for('index'))
        
    # Check 60% rule
    attendance_stats = db.execute('''
        SELECT 
            (SELECT COUNT(DISTINCT date) FROM attendance) as total_days,
            COUNT(*) as present_days
        FROM attendance 
        WHERE student_id = ? AND status = 'Present'
    ''', (student_id,)).fetchone()
    
    total_days = attendance_stats['total_days'] or 0
    present_days = attendance_stats['present_days'] or 0
    percentage = (present_days / total_days * 100) if total_days > 0 else 100
    
    has_paid_condonation = any(f['payment_type'] == 'Condonation' for f in fees)
    if percentage < 60 and not has_paid_condonation:
        flash('Attendance below 60%. Condonation fee payment is required.')
        return redirect(url_for('index'))
        
    # 3. Fetch Subjects
    subjects = db.execute('SELECT * FROM subjects WHERE class_name = ?', (student['class_name'],)).fetchall()
    
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    
    # Header
    logo_path = os.path.join(BASE_DIR, 'static', 'logo.jpg')
    if os.path.exists(logo_path):
        pdf.image(logo_path, x=95, y=10, w=20)
        pdf.ln(15)

    pdf.set_font("Times", "B", 20)
    pdf.set_text_color(26, 54, 93)
    pdf.cell(0, 12, "BHARATHIDASAN UNIVERSITY", align='C', ln=1)
    
    pdf.set_font("Times", "", 10)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 5, "Tiruchirappalli, Tamil Nadu, India", align='C', ln=1)
    pdf.cell(0, 5, "Established by the Government of Tamil Nadu", align='C', ln=1)
    
    pdf.set_draw_color(200, 152, 44)
    pdf.line(10, 40, 200, 40)
    pdf.ln(10)
    
    # Title
    pdf.set_font("Times", "B", 16)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(0, 10, "EXAMINATION HALL TICKET", align='C', ln=1)
    pdf.ln(5)
    
    # Student details
    pdf.set_font("Times", "B", 14)
    pdf.set_text_color(123, 26, 44)
    pdf.cell(0, 10, "Student Details", ln=1)
    
    pdf.set_font("Times", "", 12)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(50, 10, "Name:", border=1)
    pdf.set_font("Times", "B", 12)
    pdf.cell(100, 10, student['name'], border=1, ln=1)
    
    pdf.set_font("Times", "", 12)
    pdf.cell(50, 10, "Roll Number:", border=1)
    pdf.set_font("Times", "B", 12)
    pdf.cell(100, 10, str(student['roll_no']), border=1, ln=1)
    
    pdf.set_font("Times", "", 12)
    pdf.cell(50, 10, "Class:", border=1)
    pdf.set_font("Times", "B", 12)
    pdf.cell(100, 10, student['class_name'], border=1, ln=1)
    
    pdf.ln(10)
    
    # Subjects
    pdf.set_font("Times", "B", 14)
    pdf.set_text_color(123, 26, 44)
    pdf.cell(0, 10, "Examination Subjects", ln=1)
    
    pdf.set_font("Times", "B", 12)
    pdf.set_text_color(255, 255, 255)
    pdf.set_fill_color(26, 54, 93)
    pdf.cell(100, 10, "Subject Name", border=1, fill=True)
    pdf.cell(90, 10, "Teacher", border=1, fill=True, ln=1)
    
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Times", "", 11)
    if subjects:
        for sub in subjects:
            pdf.cell(100, 10, sub['name'], border=1)
            pdf.cell(90, 10, sub['teacher_name'] or 'N/A', border=1, ln=1)
    else:
        pdf.cell(190, 10, "No subjects registered for this class.", border=1, ln=1)
        
    pdf.ln(10)
    pdf.set_font("Times", "I", 10)
    pdf.set_text_color(185, 28, 47)
    pdf.cell(0, 10, "Note: Please bring this hall ticket to the examination hall along with your ID card.", ln=1)
    
    pdf_bytes = pdf.output()
    output = io.BytesIO(pdf_bytes)
    
    filename = f"HallTicket_{student['roll_no']}.pdf"
    return send_file(output, download_name=filename, as_attachment=True, mimetype='application/pdf')


if __name__ == '__main__':
    if not os.path.exists(DATABASE):
        init_db()
        # Create a default admin user if not exists
        with app.app_context():
            db = get_db()
            if db.execute('SELECT id FROM users').fetchone() is None:
                db.execute(
                    'INSERT INTO users (username, password_hash) VALUES (?, ?)',
                    ('admin', generate_password_hash('admin'))
                )
                db.commit()
                print("Created default user 'admin' with password 'admin'")
    
    app.run(debug=True)
